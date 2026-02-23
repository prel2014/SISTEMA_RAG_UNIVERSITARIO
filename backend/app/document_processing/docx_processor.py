import zipfile
import os

from docx import Document
from PIL import Image
import io


_IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}


class DocxProcessor:
    def extract_text(self, file_path, app=None):
        """Extract text (paragraphs + tables) from DOCX.

        When app is provided and VISION_ENABLED is true, also extracts
        embedded images from word/media/ and appends LLM descriptions.
        """
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)

        content = '\n'.join(full_text)
        if not content.strip():
            raise ValueError('No se pudo extraer texto del documento DOCX.')

        vision_enabled = app is not None and app.config.get('VISION_ENABLED', False)
        if vision_enabled:
            descriptions = self._extract_image_descriptions(file_path, app)
            if descriptions:
                content = content + '\n\n[Imágenes:\n' + '\n'.join(descriptions) + '\n]'

        return [{'content': content, 'page': 1}]

    # ------------------------------------------------------------------

    def _extract_image_descriptions(self, file_path, app):
        from app.document_processing.vision import describe_image

        min_size   = app.config.get('VISION_MIN_IMAGE_SIZE', 80)
        max_images = app.config.get('VISION_MAX_IMAGES_PAGE', 5)
        ollama_url = app.config.get('OLLAMA_BASE_URL', 'http://localhost:11434')
        llm_model  = app.config.get('LLM_MODEL', 'gemma3:4b')

        descriptions = []
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                media_entries = [
                    name for name in z.namelist()
                    if name.startswith('word/media/')
                    and os.path.splitext(name)[1].lower() in _IMAGE_EXTENSIONS
                ]
                for name in media_entries:
                    if len(descriptions) >= max_images:
                        break
                    try:
                        img_bytes = z.read(name)
                        img = Image.open(io.BytesIO(img_bytes))
                        w, h = img.size
                        if w < min_size or h < min_size:
                            continue
                        desc = describe_image(img_bytes, ollama_url, llm_model)
                        if desc:
                            idx = len(descriptions) + 1
                            descriptions.append(f'[Imagen {idx}: {desc}]')
                            print(f'[Vision] DOCX img {idx} ({name}): {desc[:80]}...')
                    except Exception as e:
                        print(f'[Vision] Error en imagen {name}: {e}')
        except Exception as e:
            print(f'[Vision] Error abriendo DOCX como ZIP: {e}')

        return descriptions
